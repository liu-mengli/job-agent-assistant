""".doc / .docx 文档解析 —— LibreOffice 转换 + python-docx 提取 + 数据清洗 + 图片提取"""

import os
import re
import shutil
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from api.log import logger
from config import settings


# --- OLE/RTF 残留控制字符 ---
_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

# --- Word 域代码残留 ---
_FIELD_CODE_RE = re.compile(
    r"\b(?:HYPERLINK|INCLUDEPICTURE|SHAPE|MERGEFORMAT|MERGEFORMATINET|TOC)\\?"
    r"[^\n]*",
    re.IGNORECASE,
)

# --- 法律声明 / 保修条款 / 环保声明 关键词（行级过滤）---
_BOILERPLATE_PATTERNS = [
    r"法律事項聲明",
    r"保\s*證\s*書",
    r"版權聲明",
    r"版權所有.*翻印必究",
    r"著作人.*致茂電子",
    r"致茂電子股份有限公司",
    r"華亞科技園區",
    r"服務專線",
    r"傳真電話",
    r"有害物質",
    r"SJ/T\s*11363",
    r"EU\s*2005/618",
    r"環保使用期限",
    r"切勿將本設備處理為未分類的廢棄物",
    r"設備及材料污染控制聲明",
    r"Chroma尚未全面完成無鉛焊錫",
    r"版本修訂紀錄",
]

# --- 孤立页码 ---
_PAGE_NUMBER_RE = re.compile(r"^\s*\d{1,3}\s*$")

# --- 多余空行 ---
_MULTI_NEWLINE_RE = re.compile(r"\n{4,}")

# --- OLE/XML 残余标签 ---
_XML_TAG_LINE_RE = re.compile(r"^\s*[</>].*[</>]\s*$")
_GARBLED_LINE_RE = re.compile(r"^[^一-鿿　-〿＀-￯a-zA-Z0-9\s\d.,;:!?()[\]{}<>@#$%^&*+=_\"'|/\\-]{10,}$")


def _clean_text(text: str) -> str:
    """数据清洗：去除 OLE 乱码、域代码、页眉页脚、法律声明白"""
    # 去除控制字符
    text = _CONTROL_CHAR_RE.sub("", text)

    # 去除 Word 域代码残留行
    text = _FIELD_CODE_RE.sub("", text)

    # 按行过滤
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()

        # 跳过空行（稍后统一处理）
        if not stripped:
            cleaned_lines.append("")
            continue

        # 跳过孤立页码
        if _PAGE_NUMBER_RE.match(stripped):
            continue

        # 跳过法律声明/保修/环保 boilerplate 行
        skip = False
        for pattern in _BOILERPLATE_PATTERNS:
            if re.search(pattern, stripped):
                skip = True
                break
        if skip:
            continue

        # 跳过过短的无意义行（OLE 残留碎片）
        if len(stripped) < 4 and not re.search(r"[a-zA-Z0-9一-鿿]", stripped):
            continue

        # 跳过 OLE/XML 残留标签行
        if _XML_TAG_LINE_RE.match(stripped):
            continue

        # 跳过乱码行（超过 10 字符且大部分不在正常字符范围内）
        if _GARBLED_LINE_RE.match(stripped):
            continue

        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)

    # 合并多余空行
    text = _MULTI_NEWLINE_RE.sub("\n\n\n", text)

    # 合并标签行与说明行：文档中大量 "Save\n\n點選後..." 结构，不合并会丢失上下文
    text = _merge_labels(text)

    return text.strip()


# 标签行特征：短行（< 50 字符），单独一行，英文字母/数字较多
_LABEL_LINE_RE = re.compile(r"^[\s]*(?:[A-Za-z0-9\[\]\(\)/ .,;:#_\-一-鿿]{1,50})$")


# 章节标题特征（与 kb_chunker 保持一致），不对其做 label merge
_SECTION_HEADER_RE = re.compile(
    r"(?:"
    r"\d+\.\d*\s+|"               # 自动编号如 "1.3 " "2. "
    r"安全概要|主畫面|主要操作介面|安全符號|"
    r"Yield Control|良率控制|"
    r"Offset\s*Setting|各軸位置修正|"
    r"Device Setting|測試設備|Tester Setup|Site Setting|"
    r"Event\s*Log|事件記錄|"
    r"Contact\s*Setting|手測與壓貨|"
    r"Tray\s*File|料盤資料|"
    r"Category\s*Setup|分類設定|"
    r"Interface\s*File|通訊介面|"
    r"Speed\s*Setting|速度設定|"
    r"Timer\s*Setting|時間設定|"
    r"Motor\s*Monitor|馬達監視|"
    r"IO\s*Monitor|IO\s*監視|"
    r"Device\s*Setup|Position\s*Check|點位確認|"
    r"Pin1\s*Check|Pin1\s*功能|"
    r"SLT\s*Test\s*Program|"
    r"Auto\s*Alignment|自動位置校正|"
    r"PM\s*Schedule|保養排程|"
    r"Run\s*Page|生產執行|"
    r"User\s*Page|作業員操作|"
    r"Setup\s*Page|設定工程師|"
    r"Engineer\s*Page|系統工程師|"
    r"使用手冊|操作手冊|"
    r"Lidar|Cobra|溫度控制|"
    r"目\s*錄|Parament|參數設定|"
    r"Lot\s*Information|生產資訊|"
    r"Output\s*Tray\s*Map|Tray\s*Map"
    r")",
    re.IGNORECASE,
)


def _merge_labels(text: str) -> str:
    """将孤立的按钮/字段标签行与其下方的说明行合并。
    原始结构:  Save\n\n點選後，將目前Offset欄位資料作儲存動作。
    合并结果:  Save: 點選後，將目前Offset欄位資料作儲存動作。
    章节标题（如 Yield Control、安全概要）不参与合并，保持独立。
    """
    lines = text.split("\n")
    merged = []
    i = 0
    while i < len(lines):
        this_line = lines[i].strip()
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
        next2 = lines[i + 2].strip() if i + 2 < len(lines) else ""

        # 章节标题不参与合并
        is_header = bool(_SECTION_HEADER_RE.match(this_line))

        # next2 是章节标题时也不合并（防止 "離開頁面。: Motor Monitor (...)"）
        next2_is_header = bool(_SECTION_HEADER_RE.match(next2))

        # 检测 pattern: label → 空行 → 说明（说明至少 10 个字符，且非章节标题）
        if (
            this_line
            and not is_header
            and _LABEL_LINE_RE.match(this_line)
            and next_line == ""
            and len(next2) >= 10
            and not next2_is_header
        ):
            merged.append(f"{this_line}: {next2}")
            i += 3  # 跳过 label、空行、说明
        else:
            if this_line:
                merged.append(this_line)
            elif merged and merged[-1] != "":
                merged.append("")
            i += 1

    return "\n".join(merged)


def _table_to_markdown(table) -> str:
    """将 python-docx Table 转为 Markdown table 格式"""
    rows = table.rows
    if not rows:
        return ""

    lines = []
    # 表头
    header_cells = [cell.text.strip().replace("\n", " ") for cell in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # 数据行
    for row in rows[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def _convert_doc_to_docx(doc_path: str, output_dir: str) -> str:
    """使用 LibreOffice headless 将 .doc 转为 .docx"""
    soffice = settings.LIBREOFFICE_PATH
    if not os.path.isfile(soffice):
        raise FileNotFoundError(
            f"LibreOffice 未找到: {soffice}\n"
            "请安装 LibreOffice 或修改 config.py 中的 LIBREOFFICE_PATH"
        )

    # 转为绝对路径，防止后台线程中相对路径失效
    abs_doc = os.path.abspath(doc_path)
    if not os.path.isfile(abs_doc):
        raise FileNotFoundError(f"源文件不存在: {abs_doc}")

    basename = os.path.splitext(os.path.basename(doc_path))[0]
    cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", output_dir, abs_doc]
    logger.info(f"执行 LibreOffice 转换: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice 转换失败 (exit={result.returncode}):\n"
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )

    expected = os.path.join(output_dir, f"{basename}.docx")
    if not os.path.isfile(expected):
        raise FileNotFoundError(f"转换后的 .docx 文件未找到: {expected}")
    return expected


def _raw_extract_doc_text(doc_path: str) -> str:
    """回退方案：直接读取 .doc 二进制文件中的 UTF-16LE 文本（无需 LibreOffice）。
    适用于旧版 OLE .doc 格式，结果质量不如 LibreOffice 转换但可应急使用。"""
    logger.warning(f"LibreOffice 不可用，使用 UTF-16LE 原始提取: {doc_path}")
    with open(doc_path, "rb") as f:
        raw = f.read()
    text = raw.decode("utf-16-le", errors="ignore")

    # 过滤出可读文本段
    current = []
    lines = []
    for ch in text:
        if ch.isprintable() or ch in "\n\r\t":
            current.append(ch)
        else:
            s = "".join(current).strip()
            if len(s) > 40:
                lines.append(s)
            current = []
    s = "".join(current).strip()
    if len(s) > 40:
        lines.append(s)

    logger.info(f"UTF-16LE 原始提取: {len(lines)} 个段落")
    return "\n\n".join(lines)


def _extract_docx(docx_path: str) -> str:
    """从 .docx 提取段落文本和表格为 Markdown，Heading 自动编号"""
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(docx_path)
    parts = []

    # 章节计数器
    h1 = 0
    h2 = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name == "Heading 1":
                h1 += 1
                h2 = 0
                parts.append(f"{h1}. {text}")
            elif style_name == "Heading 2":
                h2 += 1
                parts.append(f"{h1}.{h2} {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            table = Table(element, doc)
            md_table = _table_to_markdown(table)
            if md_table:
                parts.append("\n" + md_table + "\n")

    return "\n\n".join(parts)


# --- 图片提取 ---
_MIN_IMAGE_SIZE = 2048  # 跳过 < 2KB 的小图标/装饰图
_SKIP_EXTENSIONS = {".wmf"}  # 浏览器不支持的矢量格式


def extract_docx_images(docx_path: str, output_dir: str) -> dict[str, str]:
    """从 docx zip 中提取所有图片到 output_dir，返回 {rId: filename} 映射（仅已提取的）"""
    os.makedirs(output_dir, exist_ok=True)
    rid_map: dict[str, str] = {}
    extracted_filenames: set[str] = set()

    with zipfile.ZipFile(docx_path) as zf:
        # 提取 media 目录下的图片文件（先提取再建映射，确保映射只含实际存在的文件）
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue

            filename = os.path.basename(name)
            ext = os.path.splitext(filename)[1].lower()

            if ext in _SKIP_EXTENSIONS:
                continue

            info = zf.getinfo(name)
            if info.file_size < _MIN_IMAGE_SIZE:
                continue

            dest = os.path.join(output_dir, filename)
            with zf.open(name) as src, open(dest, "wb") as dst:
                shutil.copyfileobj(src, dst)
            extracted_filenames.add(filename)

        # 解析关系文件，构建 rId → 图片文件的映射（仅包含已提取的文件）
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in zf.namelist():
            rels_xml = zf.read(rels_path)
            ns = "{http://schemas.openxmlformats.org/package/2006/relationships}"
            root = ET.fromstring(rels_xml)
            for rel in root:
                rtype = rel.get("Type", "")
                if "image" not in rtype.lower():
                    continue
                rid = rel.get("Id")
                target = rel.get("Target")
                if rid and target:
                    filename = os.path.basename(target)
                    if filename in extracted_filenames:
                        rid_map[rid] = filename

    logger.info(f"图片提取完成: {len(extracted_filenames)} 张 → {output_dir}，{len(rid_map)} 个 rId 映射")
    return rid_map


def map_images_to_sections(
    docx_path: str, image_rid_map: dict[str, str]
) -> dict[str, list[str]]:
    """遍历 docx 正文，将图片按 Heading 章节归属，返回 {section_label: [filename, ...]}

    section_label 格式与 kb_chunker.py 的 _is_section_header 保持一致：
    - Heading 1 → "1. 标题"、"4. Engineer Page"
    - Heading 2 → "4.10 Position Check"
    """
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    A_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

    doc = Document(docx_path)
    h1, h2 = 0, 0
    current_section = "全文"
    section_images: dict[str, list[str]] = {}

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            style_name = para.style.name if para.style else ""

            # 检测章节标题切换
            if style_name == "Heading 1" and text:
                h1 += 1
                h2 = 0
                current_section = f"{h1}. {text}"
            elif style_name == "Heading 2" and text:
                h2 += 1
                current_section = f"{h1}.{h2} {text}"

            # 检测图片
            blips = para._element.findall(f".//{A_BLIP}")
            for blip in blips:
                rid = blip.get(R_EMBED)
                if rid and rid in image_rid_map:
                    filename = image_rid_map[rid]
                    section_images.setdefault(current_section, []).append(filename)

        elif tag == "tbl":
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        blips = para._element.findall(f".//{A_BLIP}")
                        for blip in blips:
                            rid = blip.get(R_EMBED)
                            if rid and rid in image_rid_map:
                                filename = image_rid_map[rid]
                                section_images.setdefault(current_section, []).append(filename)

    # 去重：同一章节内图片可能被多次引用
    for section in section_images:
        seen: list[str] = []
        deduped = []
        for f in section_images[section]:
            if f not in seen:
                seen.append(f)
                deduped.append(f)
        section_images[section] = deduped

    total = sum(len(v) for v in section_images.values())
    logger.info(f"图片章节映射完成: {total} 张 → {len(section_images)} 个章节")
    return section_images


def parse_doc_with_images(
    file_path: str, images_output_dir: str
) -> tuple[str, dict[str, list[str]]]:
    """解析 .doc/.docx 文件，同时提取图片。

    Returns:
        (cleaned_text, section_images) — section_images 的 key 与 kb_chunker 章节标签一致
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".doc":
        soffice = settings.LIBREOFFICE_PATH
        if os.path.isfile(soffice):
            logger.info(f"开始转换 .doc 文件: {file_path}")
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = _convert_doc_to_docx(file_path, tmpdir)
                logger.info(f".doc → .docx 转换完成: {docx_path}")
                return _parse_docx_stage(docx_path, images_output_dir)
        else:
            raw_text = _raw_extract_doc_text(file_path)
            cleaned = _clean_text(raw_text)
            logger.warning("LibreOffice 不可用，图片提取不可用")
            return cleaned, {}
    elif ext == ".docx":
        logger.info(f"开始解析 .docx 文件: {file_path}")
        return _parse_docx_stage(file_path, images_output_dir)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .doc / .docx")


def _parse_docx_stage(
    docx_path: str, images_output_dir: str
) -> tuple[str, dict[str, list[str]]]:
    """内部：从 docx 同时提取文本和图片"""
    rid_map = extract_docx_images(docx_path, images_output_dir)
    section_images = map_images_to_sections(docx_path, rid_map) if rid_map else {}
    raw_text = _extract_docx(docx_path)
    cleaned = _clean_text(raw_text)
    logger.info(f"DOC 解析完成: {len(raw_text)} → {len(cleaned)} 字符（清洗后）")
    return cleaned, section_images


def parse_doc(file_path: str) -> str:
    """解析 .doc 或 .docx 文件，返回清洗后的纯文本（含 Markdown 表格）"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".doc":
        soffice = settings.LIBREOFFICE_PATH
        if os.path.isfile(soffice):
            logger.info(f"开始转换 .doc 文件: {file_path}")
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = _convert_doc_to_docx(file_path, tmpdir)
                logger.info(f".doc → .docx 转换完成: {docx_path}")
                raw_text = _extract_docx(docx_path)
        else:
            # LibreOffice 不可用时回退到原始二进制提取
            raw_text = _raw_extract_doc_text(file_path)
    elif ext == ".docx":
        logger.info(f"开始解析 .docx 文件: {file_path}")
        raw_text = _extract_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .doc / .docx")

    cleaned = _clean_text(raw_text)
    logger.info(f"DOC 解析完成: {len(raw_text)} → {len(cleaned)} 字符（清洗后）")
    return cleaned
